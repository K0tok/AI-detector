#!/usr/bin/env python3
"""
Improved AI Text Detector for DOCX Files (Russian)
This program helps users identify AI-generated text in DOCX documents
and suggests improvements to make it appear more human-written.
"""

import os
import sys
import re
import docx
from docx import Document
from collections import Counter
import argparse
import json
from typing import List, Tuple, Dict
import statistics


class AITextDetector:
    """
    An improved class to detect potentially AI-generated text in DOCX documents
    and provide suggestions for making it more human-like.
    """
    
    def __init__(self):
        # Russian-specific patterns that might indicate AI-generated text
        self.ai_indicators = {
            'overly_formal': [
                r'\\b(?:использовать|применять|осуществлять|реализовывать|обеспечивать|содействовать|способствовать|предоставлять|представлять|демонстрировать|иллюстрировать|осуществлять|последовательно|однако|более того|кроме того|следовательно|несмотря на|в соответствии с|в котором|в то время как|в связи с|в рамках|в целях)\\b',
            ],
            'repetitive_transitions': [
                r'(?:во-первых|во-вторых|в-третьих),',
                r'(?:более того|кроме того|также|в дополнение|следовательно|поэтому|в результате|по этой причине|в связи с этим|в заключение|в итоге|в качестве того|в свою очередь|с другой стороны)',
            ],
            'generic_structure': [
                r'(?:в современном мире|в настоящее время|в условиях|в рамках|в контексте|с учетом|в свете|в связи с вышеизложенным)',
            ],
            'formulaic_phrases': [
                r'(?:с другой стороны|в заключение|в итоге|в общем и целом|на сегодняшний день|в настоящее время|в условиях современности|как уже упоминалось|как было сказано|в заключение стоит отметить)',
            ],
            'repetitive_start': [
                r'^\\s*в\\s',  # Starting too many sentences with "в" (In)
            ]
        }
        
        # Advanced patterns for more sophisticated detection
        self.advanced_patterns = {
            'sentence_complexity': r'[.!?].*?[,;].*?[,;].*?[.!?]',  # Sentences with multiple commas/semicolons
            'passive_voice': r'\\b(?:был|была|было|были|будет|будут|является|являются|представляет|представляют|состоит|составляет)\\b.*?[,;].*?\\b(?:что|который|которая|которое)\\b',
            'repetitive_punctuation': r'[,;]{2,}',  # Multiple commas or semicolons
            'unnatural_punctuation': r'[.!?] [,;]',  # Punctuation followed by comma/semicolon
        }
        
        # N-gram patterns that are common in AI text
        self.ngram_patterns = [
            'в связи с вышеизложенным',
            'в заключение стоит отметить',
            'следует отметить, что',
            'в целях обеспечения',
            'в рамках реализации',
            'в условиях современности',
            'в свете вышеизложенного',
            'в контексте современных',
            'с учетом вышеизложенного',
            'в качестве основного',
            'в качестве важного',
            'в качестве ключевого',
        ]
        
        # Suggestions for improving human-like writing in Russian
        self.suggestions = {
            'overly_formal': {
                'осуществлять': 'делать',
                'реализовывать': 'выполнять',
                'обеспечивать': 'обеспечивать',
                'содействовать': 'помогать',
                'способствовать': 'помогать',
                'предоставлять': 'давать',
                'представлять': 'показывать',
                'демонстрировать': 'показывать',
                'иллюстрировать': 'показывать',
                'использовать': 'использовать',
                'применять': 'использовать',
            },
            'generic_transitions': {
                'более того': 'и',
                'кроме того': 'и еще',
                'следовательно': 'поэтому',
                'в результате': 'так что',
                'поэтому': 'так что',
                'в связи с этим': 'из-за этого',
                'в заключение': 'в конце',
                'в итоге': 'в результате',
                'в общем и целом': 'в общем',
            }
        }

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the document
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Ошибка при чтении DOCX файла: {e}")
            return ""

    def calculate_sentence_stats(self, text: str) -> Dict:
        """
        Calculate various sentence statistics for AI detection.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with sentence statistics
        """
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        
        if not sentences:
            return {}
        
        # Calculate average sentence length
        sentence_lengths = [len(s.split()) for s in sentences if s]
        avg_sentence_length = statistics.mean(sentence_lengths) if sentence_lengths else 0
        
        # Calculate sentence length variance (AI tends to have more uniform lengths)
        sentence_variance = statistics.variance(sentence_lengths) if len(sentence_lengths) > 1 else 0
        
        # Check for repetitive sentence starts
        first_words = [s.split()[0].lower() if s.split() else '' for s in sentences]
        first_word_counts = Counter(first_words)
        most_common_first_word = first_word_counts.most_common(1)[0] if first_word_counts else ('', 0)
        
        # Check for sentence complexity (AI often uses complex sentences)
        complex_sentences = len([s for s in sentences if len(re.findall(r'[,;]', s)) >= 2])
        complex_sentence_ratio = complex_sentences / len(sentences) if sentences else 0
        
        return {
            'avg_sentence_length': avg_sentence_length,
            'sentence_variance': sentence_variance,
            'most_common_first_word': most_common_first_word,
            'complex_sentence_ratio': complex_sentence_ratio,
            'total_sentences': len(sentences)
        }

    def calculate_word_stats(self, text: str) -> Dict:
        """
        Calculate various word statistics for AI detection.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with word statistics
        """
        words = [w.lower() for w in re.findall(r'\b\w+\b', text)]
        
        if not words:
            return {}
        
        # Calculate average word length
        word_lengths = [len(w) for w in words]
        avg_word_length = statistics.mean(word_lengths) if word_lengths else 0
        
        # Calculate word frequency distribution (AI text often has more uniform distribution)
        word_freq = Counter(words)
        unique_words = len(word_freq)
        total_words = len(words)
        uniqueness_ratio = unique_words / total_words if total_words > 0 else 0
        
        # Calculate average frequency of words (lower = more varied vocabulary)
        avg_word_freq = statistics.mean(word_freq.values()) if word_freq.values() else 0
        
        # Check for formal vs informal word ratio
        formal_words = sum(1 for w in words if w in ['осуществлять', 'реализовывать', 'обеспечивать', 'содействовать', 'способствовать', 'предоставлять', 'представлять'])
        formal_ratio = formal_words / total_words if total_words > 0 else 0
        
        return {
            'avg_word_length': avg_word_length,
            'unique_words': unique_words,
            'total_words': total_words,
            'uniqueness_ratio': uniqueness_ratio,
            'avg_word_freq': avg_word_freq,
            'formal_ratio': formal_ratio
        }

    def analyze_text(self, text: str) -> Dict:
        """
        Analyze text for AI-generated patterns using multiple methods.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with analysis results
        """
        results = {
            'total_chars': len(text),
            'total_words': len(text.split()),
            'total_sentences': len(re.split(r'[.!?]+', text)),
            'issues': [],
            'ai_probability': 0,
            'suggestions': [],
            'sentence_stats': {},
            'word_stats': {},
            'detailed_analysis': {}
        }
        
        # Calculate statistics
        results['sentence_stats'] = self.calculate_sentence_stats(text)
        results['word_stats'] = self.calculate_word_stats(text)
        
        # Calculate initial probability based on statistics
        prob_from_stats = self.calculate_probability_from_stats(results['sentence_stats'], results['word_stats'])
        
        # Check for each pattern category
        for category, patterns in self.ai_indicators.items():
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                if matches:
                    count = len(matches)
                    results['issues'].append({
                        'type': category,
                        'pattern': pattern,
                        'count': count,
                        'examples': matches[:5]  # Show first 5 examples
                    })
                    
                    # Increase AI probability based on findings
                    if category == 'overly_formal':
                        results['ai_probability'] += count * 3
                    elif category == 'repetitive_transitions':
                        results['ai_probability'] += count * 4
                    elif category == 'generic_structure':
                        results['ai_probability'] += count * 5
                    elif category == 'formulaic_phrases':
                        results['ai_probability'] += count * 4
                    elif category == 'repetitive_start':
                        results['ai_probability'] += count * 2
        
        # Check for advanced patterns
        for pattern_name, pattern in self.advanced_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            if matches:
                count = len(matches)
                results['issues'].append({
                    'type': pattern_name,
                    'pattern': pattern,
                    'count': count,
                    'examples': matches[:3]
                })
                
                if pattern_name == 'sentence_complexity':
                    results['ai_probability'] += count * 2
                elif pattern_name == 'passive_voice':
                    results['ai_probability'] += count * 3
                elif pattern_name == 'repetitive_punctuation':
                    results['ai_probability'] += count * 5
                elif pattern_name == 'unnatural_punctuation':
                    results['ai_probability'] += count * 4
        
        # Check for n-gram patterns
        ngram_matches = []
        for ngram in self.ngram_patterns:
            found = re.findall(re.escape(ngram), text, re.IGNORECASE)
            if found:
                ngram_matches.extend(found)
        
        if ngram_matches:
            results['issues'].append({
                'type': 'common_ngrams',
                'pattern': 'N-gram patterns',
                'count': len(ngram_matches),
                'examples': ngram_matches[:5]
            })
            results['ai_probability'] += len(ngram_matches) * 6
        
        # Add statistical probability
        results['ai_probability'] += prob_from_stats
        
        # Normalize probability to 0-100 range
        results['ai_probability'] = min(max(results['ai_probability'], 0), 100)
        
        # Store detailed analysis
        results['detailed_analysis'] = {
            'statistical_probability': prob_from_stats,
            'pattern_probability': results['ai_probability'] - prob_from_stats if results['ai_probability'] > prob_from_stats else 0
        }
        
        return results

    def calculate_probability_from_stats(self, sentence_stats: Dict, word_stats: Dict) -> float:
        """
        Calculate AI probability based on statistical analysis.
        
        Args:
            sentence_stats: Sentence statistics
            word_stats: Word statistics
            
        Returns:
            Probability score (0-100)
        """
        score = 0.0
        
        # Sentence analysis
        if sentence_stats:
            # Low variance in sentence length suggests AI
            if sentence_stats.get('sentence_variance', 0) < 20:  # Lower than typical human variance
                score += 15
            
            # High ratio of complex sentences suggests AI
            if sentence_stats.get('complex_sentence_ratio', 0) > 0.3:  # More than 30% complex sentences
                score += 20
            
            # Repetitive sentence starts
            if sentence_stats.get('most_common_first_word', ('', 0))[1] > sentence_stats.get('total_sentences', 1) * 0.2:  # More than 20% start the same
                score += 25
        
        # Word analysis
        if word_stats:
            # Low uniqueness ratio suggests AI
            if word_stats.get('uniqueness_ratio', 1) < 0.5:  # Less than 50% unique words
                score += 15
            
            # High formal word ratio suggests AI
            if word_stats.get('formal_ratio', 0) > 0.05:  # More than 5% formal words
                score += 20
        
        return min(score, 100)

    def generate_suggestions(self, text: str, analysis: Dict = None) -> List[Dict]:
        """
        Generate suggestions to make text appear more human-written.
        
        Args:
            text: Text to analyze
            analysis: Pre-computed analysis results (optional)
            
        Returns:
            List of suggestions
        """
        suggestions = []
        
        # Check for overly formal words and suggest alternatives
        for word, replacement in self.suggestions['overly_formal'].items():
            if re.search(r'\\b' + re.escape(word) + r'\\b', text, re.IGNORECASE):
                suggestions.append({
                    'issue': f'Слишком формальное слово: "{word}"',
                    'suggestion': f'Замените "{word}" на более простые альтернативы вроде "{replacement}"',
                    'type': 'word_replacement'
                })
        
        # Check for repetitive sentence starters
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        v_starts = [s for s in sentences if s.strip().lower().startswith('в ')]
        if len(v_starts) > len(sentences) * 0.3:  # If more than 30% start with "в"
            suggestions.append({
                'issue': f'Слишком много предложений начинаются с "в" ({len(v_starts)}/{len(sentences)})',
                'suggestion': 'Варьируйте начало предложений, чтобы избежать повторяющихся шаблонов',
                'type': 'sentence_structure'
            })
        
        # Check for repetitive transitions
        transition_matches = []
        for pattern in self.ai_indicators['repetitive_transitions']:
            transition_matches.extend(re.findall(pattern, text, re.IGNORECASE))
        
        if len(transition_matches) > len(sentences) * 0.1:  # If more than 10% use these transitions
            suggestions.append({
                'issue': f'Слишком много повторяющихся переходов вроде "{transition_matches[0] if transition_matches else 'transition'}"',
                'suggestion': 'Используйте разнообразные слова перехода или удаляйте ненужные',
                'type': 'transitions'
            })
        
        # Check for complex sentence structures
        if analysis and analysis.get('sentence_stats', {}).get('complex_sentence_ratio', 0) > 0.3:
            suggestions.append({
                'issue': 'Слишком много сложных предложений с несколькими запятыми/точками с запятой',
                'suggestion': 'Разбейте сложные предложения на более простые, чтобы текст звучал естественнее',
                'type': 'sentence_structure'
            })
        
        # Check for low vocabulary diversity
        if analysis and analysis.get('word_stats', {}).get('uniqueness_ratio', 1) < 0.5:
            suggestions.append({
                'issue': f'Низкое разнообразие словарного запаса (только {analysis["word_stats"].get("uniqueness_ratio", 0)*100:.1f}% уникальных слов)',
                'suggestion': 'Используйте более разнообразный словарный запас, чтобы избежать повторений',
                'type': 'vocabulary'
            })
        
        return suggestions

    def process_docx_file(self, file_path: str) -> Dict:
        """
        Process a DOCX file and return analysis results.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Analysis results
        """
        if not os.path.exists(file_path):
            return {'error': f'Файл не существует: {file_path}'}
        
        if not file_path.lower().endswith('.docx'):
            return {'error': 'Файл должен быть DOCX документом'}
        
        # Extract text from the document
        text = self.extract_text_from_docx(file_path)
        
        if not text.strip():
            return {'error': 'В документе не найден текст'}
        
        # Analyze the text
        analysis = self.analyze_text(text)
        suggestions = self.generate_suggestions(text, analysis)
        
        # Combine results
        result = {
            'file_path': file_path,
            'text_sample': text[:200] + '...' if len(text) > 200 else text,
            'analysis': analysis,
            'suggestions': suggestions,
            'summary': {
                'ai_probability': analysis['ai_probability'],
                'issue_count': len(analysis['issues']),
                'suggestion_count': len(suggestions),
                'text_length': analysis['total_chars']
            }
        }
        
        return result

    def fix_text(self, text: str) -> str:
        """
        Apply fixes to make text appear more human-written.
        
        Args:
            text: Original text
            
        Returns:
            Improved text
        """
        fixed_text = text
        
        # Replace overly formal words with simpler alternatives
        for old_word, new_word in self.suggestions['overly_formal'].items():
            # Use word boundaries to match whole words only
            fixed_text = re.sub(r'\\b' + re.escape(old_word) + r'\\b', new_word, fixed_text, flags=re.IGNORECASE)
        
        # Additional Russian-specific replacements for common AI phrases
        russian_replacements = {
            r'\\bв современном мире\\b': 'сейчас',
            r'\\bв настоящее время\\b': 'сейчас',
            r'\\bследует отметить, что\\b': 'отметим, что',
            r'\\bв целях обеспечения\\b': 'чтобы обеспечить',
            r'\\bв рамках реализации\\b': 'при выполнении',
            r'\\bв условиях\\b': 'при',
            r'\\bв контексте\\b': 'в связи с',
            r'\\bс учетом\\b': 'учитывая',
            r'\\bв свете\\b': 'из-за',
            r'\\bв связи с вышеизложенным\\b': 'поэтому',
            r'\\bкак уже упоминалось\\b': 'как говорилось',
            r'\\bкак было сказано\\b': 'как говорили',
            r'\\bв заключение стоит отметить\\b': 'в заключение',
            r'\\bв итоге\\b': 'в результате',
            r'\\bв общем и целом\\b': 'в общем',
            r'\\bна сегодняшний день\\b': 'сейчас',
            r'\\bв условиях современности\\b': 'сейчас',
            r'\\bв свете вышеизложенного\\b': 'поэтому',
            r'\\bс другой стороны\\b': 'а еще',
            r'\\bв заключение\\b': 'в конце',
            r'\\bв качестве основного\\b': 'как основу',
            r'\\bв качестве важного\\b': 'важно',
            r'\\bв качестве ключевого\\b': 'ключевым',
            r'\\bболее того\\b': 'и еще',
            r'\\bкроме того\\b': 'и',
            r'\\bтакже\\b': 'и',
            r'\\bв дополнение\\b': 'еще',
            r'\\bследовательно\\b': 'поэтому',
            r'\\bпоэтому\\b': 'так что',
            r'\\bв результате\\b': 'так что',
            r'\\bпо этой причине\\b': 'поэтому',
            r'\\bв связи с этим\\b': 'поэтому',
            r'\\bво-первых\\b': 'первое',
            r'\\bво-вторых\\b': 'второе',
            r'\\bв-третьих\\b': 'третье',
        }
        
        for old_pattern, new_replacement in russian_replacements.items():
            fixed_text = re.sub(old_pattern, new_replacement, fixed_text, flags=re.IGNORECASE)
        
        # Add more natural variations to repetitive sentence structures
        # Replace repetitive "в [context], [statement]" patterns
        fixed_text = re.sub(
            r'\\bв\\s+(\\w+)\\s*,\\s*([а-яёА-ЯЁ])', 
            r'\\1 — \\2', 
            fixed_text, 
            flags=re.IGNORECASE
        )
        
        # Try to vary sentence beginnings
        fixed_text = re.sub(r'(^|\\.\\s+)в\\s+', r'\\1', fixed_text, count=0, flags=re.MULTILINE)
        
        return fixed_text


def save_report(results: Dict, output_path: str):
    """
    Save analysis results to a file.
    
    Args:
        results: Analysis results
        output_path: Path to save the report
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"Отчет сохранен в: {output_path}")


def print_detailed_results(results: Dict):
    """
    Print detailed analysis results to console.
    
    Args:
        results: Analysis results
    """
    print("\n" + "="*60)
    print("РЕЗУЛЬТАТЫ ОБНАРУЖЕНИЯ ТЕКСТА ИИ")
    print("="*60)
    print(f"Файл: {results['file_path']}")
    print(f"Длина текста: {results['analysis']['total_chars']} символов, {results['analysis']['total_words']} слов")
    print(f"Оценка вероятности ИИ: {results['analysis']['ai_probability']:.1f}%")
    
    # Show detailed statistics
    sentence_stats = results['analysis'].get('sentence_stats', {})
    word_stats = results['analysis'].get('word_stats', {})
    
    if sentence_stats:
        print(f"Средняя длина предложения: {sentence_stats.get('avg_sentence_length', 0):.1f} слов")
        print(f"Доля сложных предложений: {sentence_stats.get('complex_sentence_ratio', 0)*100:.1f}%")
    
    if word_stats:
        print(f"Уникальные слова: {word_stats.get('uniqueness_ratio', 0)*100:.1f}%")
        print(f"Средняя длина слова: {word_stats.get('avg_word_length', 0):.1f} символов")
    
    if results['analysis']['issues']:
        print(f"\nНайденные проблемы ({len(results['analysis']['issues'])}):")
        for issue in results['analysis']['issues']:
            print(f"  - {issue['type']}: {issue['count']} вхождений")
            if issue.get('examples'):
                print(f"    Примеры: {', '.join(issue['examples'][:3])}")
    else:
        print("\nКрупных шаблонов ИИ не обнаружено.")
    
    if results['suggestions']:
        print(f"\nПредложения ({len(results['suggestions'])}):")
        for i, suggestion in enumerate(results['suggestions'], 1):
            print(f"  {i}. {suggestion['issue']}")
            print(f"     {suggestion['suggestion']}")
    else:
        print("\nНет предложений по улучшению.")
    
    print("\n" + "="*60)


def main():
    parser = argparse.ArgumentParser(
        description='Обнаружение текста, сгенерированного ИИ, в DOCX файлах и предложения по улучшению.'
    )
    parser.add_argument('input_file', help='Путь к входному DOCX файлу')
    parser.add_argument('-o', '--output', help='Путь для сохранения отчета анализа (формат JSON)')
    parser.add_argument('--fix', action='store_true', help='Создать исправленную версию документа')
    parser.add_argument('--fix-output', help='Путь для сохранения исправленного документа (требуется --fix)')
    parser.add_argument('--detailed', action='store_true', help='Показать подробную статистику')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_file):
        print(f"Ошибка: Файл не существует: {args.input_file}")
        sys.exit(1)
    
    detector = AITextDetector()
    
    # Process the document
    results = detector.process_docx_file(args.input_file)
    
    if 'error' in results:
        print(f"Ошибка: {results['error']}")
        sys.exit(1)
    
    # Print results
    if args.detailed:
        print_detailed_results(results)
    else:
        print("\n" + "="*60)
        print("РЕЗУЛЬТАТЫ ОБНАРУЖЕНИЯ ТЕКСТА ИИ")
        print("="*60)
        print(f"Файл: {results['file_path']}")
        print(f"Длина текста: {results['analysis']['total_chars']} символов, {results['analysis']['total_words']} слов")
        print(f"Оценка вероятности ИИ: {results['analysis']['ai_probability']:.1f}%")
        
        if results['analysis']['issues']:
            print(f"\nНайдено проблем: {len(results['analysis']['issues'])}")
        else:
            print("\nКрупных шаблонов ИИ не обнаружено.")
        
        if results['suggestions']:
            print(f"Предложений по улучшению: {len(results['suggestions'])}")
        else:
            print("Нет предложений по улучшению.")
        
        print("="*60)
    
    # Save report if requested
    if args.output:
        save_report(results, args.output)
    
    # Generate fixed document if requested
    if args.fix:
        if not args.fix_output:
            print("Ошибка: --fix-output требуется при использовании --fix")
            sys.exit(1)
        
        # Extract text and apply fixes
        original_doc = Document(args.input_file)
        detector = AITextDetector()
        
        # Process paragraphs
        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():
                paragraph.text = detector.fix_text(paragraph.text)
        
        # Process tables
        for table in original_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = detector.fix_text(cell.text)
        
        # Save the fixed document
        original_doc.save(args.fix_output)
        print(f"Исправленный документ сохранен в: {args.fix_output}")


if __name__ == "__main__":
    main()