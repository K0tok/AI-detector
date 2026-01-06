#!/usr/bin/env python3
"""
AI Text Detector for DOCX Files (Russian)
This program helps users identify AI-generated text in DOCX documents
and suggests improvements to make it appear more human-written.
"""

import os
import sys
import re
import docx
from docx import Document
from collections import Counter
import argparse
import json
from typing import List, Tuple, Dict


class AITextDetector:
    """
    A class to detect potentially AI-generated text in DOCX documents
    and provide suggestions for making it more human-like.
    """
    
    def __init__(self):
        # Russian-specific patterns that might indicate AI-generated text
        self.ai_indicators = [
            # Overly formal language in Russian
            r'\b(?:использовать|применять|осуществлять|реализовывать|обеспечивать|содействовать|способствовать|предоставлять|представлять|демонстрировать|иллюстрировать|последовательно|однако|более того|кроме того|следовательно|несмотря на|в соответствии с|до настоящего времени|в котором|в то время как)\b',
            # Repetitive sentence structures
            r'(?:Во-первых|Во-вторых|В-третьих),',
            # Generic transitions in Russian
            r'(?:Более того|Кроме того|Также|В дополнение|Следовательно|В результате|Поэтому|Таким образом|Однако|Тем не менее|С другой стороны)',
            # Overuse of complex words in Russian
            r'\b(?:осуществлять|реализовывать|обеспечивать|содействовать|способствовать|демонстрировать|иллюстрировать|осуществлять)\b',
            # Repetitive phrases in Russian
            r'(?:с другой стороны|в заключение|в итоге|в свете вышеизложенного|следует отметить|что касается|в связи с этим)',
        ]
        
        # Patterns for repetitive sentence beginnings
        self.repetitive_start_patterns = [
            r'^\s*В\s',  # Starting too many sentences with "В" (In)
        ]
        
        # Common AI writing patterns in Russian
        self.patterns = {
            'repetitive_transitions': r'(?:Более того|Кроме того|Также|В дополнение|Следовательно|В результате|Поэтому|Таким образом|Однако|Тем не менее|С другой стороны)',
            'overly_formal': r'\b(?:использовать|применять|осуществлять|реализовывать|обеспечивать|содействовать|способствовать|предоставлять|представлять|демонстрировать|иллюстрировать|последовательно|однако|более того|кроме того|следовательно|несмотря на|в соответствии с|до настоящего времени|в котором|в то время как)\b',
            'generic_structure': r'(?:Во-первых|Во-вторых|В-третьих),',
            'complex_word_usage': r'\b(?:осуществлять|реализовывать|обеспечивать|содействовать|способствовать|демонстрировать|иллюстрировать)\b',
            'formulaic_phrases': r'(?:с другой стороны|в заключение|в итоге|в свете вышеизложенного|следует отметить|что касается|в связи с этим|в современном мире|в условиях|в рамках|на сегодняшний день|в настоящее время)',
            'repetitive_start': r'^\s*В\s',
        }
        
        # Suggestions for improving human-like writing in Russian
        self.suggestions = {
            'overly_formal': {
                'осуществлять': 'делать',
                'реализовывать': 'выполнять',
                'обеспечивать': 'обеспечивать',
                'содействовать': 'помогать',
                'способствовать': 'помогать',
                'демонстрировать': 'показывать',
                'иллюстрировать': 'показывать',
                'предоставлять': 'давать',
                'представлять': 'показывать',
            }
        }

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the document
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Ошибка при чтении DOCX файла: {e}")
            return ""

    def analyze_text(self, text: str) -> Dict:
        """
        Analyze text for AI-generated patterns.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with analysis results
        """
        results = {
            'total_chars': len(text),
            'total_words': len(text.split()),
            'total_sentences': len(re.split(r'[.!?]+', text)),
            'issues': [],
            'ai_probability': 0,
            'suggestions': []
        }
        
        # Check for each pattern
        for pattern_name, pattern in self.patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            if matches:
                count = len(matches)
                results['issues'].append({
                    'type': pattern_name,
                    'pattern': pattern,
                    'count': count,
                    'examples': matches[:5]  # Show first 5 examples
                })
                
                # Increase AI probability based on findings
                if pattern_name == 'overly_formal':
                    results['ai_probability'] += count * 5
                elif pattern_name == 'repetitive_transitions':
                    results['ai_probability'] += count * 8
                elif pattern_name == 'generic_structure':
                    results['ai_probability'] += count * 10
                elif pattern_name == 'formulaic_phrases':
                    results['ai_probability'] += count * 7
        
        # Calculate final probability (capped at 100)
        results['ai_probability'] = min(results['ai_probability'], 100)
        
        return results

    def generate_suggestions(self, text: str) -> List[Dict]:
        """
        Generate suggestions to make text appear more human-written.
        
        Args:
            text: Text to analyze
            
        Returns:
            List of suggestions
        """
        suggestions = []
        
        # Check for overly formal words and suggest alternatives
        for word, replacement in self.suggestions['overly_formal'].items():
            if re.search(r'\b' + re.escape(word) + r'\b', text, re.IGNORECASE):
                suggestions.append({
                    'issue': f'Слишком формальное слово: "{word}"',
                    'suggestion': f'Замените "{word}" на более простые альтернативы вроде "{replacement}"',
                    'type': 'word_replacement'
                })
        
        # Check for repetitive sentence starters
        sentences = re.split(r'[.!?]+', text)
        v_starts = [s for s in sentences if s.strip().lower().startswith('в ')]
        if len(v_starts) > len(sentences) * 0.3:  # If more than 30% start with "В"
            suggestions.append({
                'issue': f'Слишком много предложений начинаются с "В" ({len(v_starts)}/{len(sentences)})',
                'suggestion': 'Варьируйте начало предложений, чтобы избежать повторяющихся шаблонов',
                'type': 'sentence_structure'
            })
        
        # Check for repetitive transitions
        transition_matches = re.findall(self.patterns['repetitive_transitions'], text, re.IGNORECASE)
        if len(transition_matches) > len(sentences) * 0.1:  # If more than 10% use these transitions
            suggestions.append({
                'issue': f'Слишком много повторяющихся переходов вроде "{transition_matches[0]}"',
                'suggestion': 'Используйте разнообразные слова перехода или удаляйте ненужные',
                'type': 'transitions'
            })
        
        return suggestions

    def process_docx_file(self, file_path: str) -> Dict:
        """
        Process a DOCX file and return analysis results.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Analysis results
        """
        if not os.path.exists(file_path):
            return {'error': f'Файл не существует: {file_path}'}
        
        if not file_path.lower().endswith('.docx'):
            return {'error': 'Файл должен быть DOCX документом'}
        
        # Extract text from the document
        text = self.extract_text_from_docx(file_path)
        
        if not text.strip():
            return {'error': 'В документе не найден текст'}
        
        # Analyze the text
        analysis = self.analyze_text(text)
        suggestions = self.generate_suggestions(text)
        
        # Combine results
        result = {
            'file_path': file_path,
            'text_sample': text[:200] + '...' if len(text) > 200 else text,
            'analysis': analysis,
            'suggestions': suggestions
        }
        
        return result

    def fix_text(self, text: str) -> str:
        """
        Apply fixes to make text appear more human-written.
        
        Args:
            text: Original text
            
        Returns:
            Improved text
        """
        fixed_text = text
        
        # Replace overly formal words with simpler alternatives
        for old_word, new_word in self.suggestions['overly_formal'].items():
            # Use word boundaries to match whole words only
            fixed_text = re.sub(r'\b' + re.escape(old_word) + r'\b', new_word, fixed_text, flags=re.IGNORECASE)
        
        # Additional Russian-specific replacements for common AI phrases
        russian_replacements = {
            r'\bВ современном мире\b': 'Сейчас',
            r'\bСледует отметить, что\b': 'Отметим, что',
            r'\bКлючевым аспектом\b': 'Важным моментом',
            r'\bСледует упомянуть, что\b': 'Упомянем, что',
            r'\bНеобходимо учитывать, что\b': 'Помним, что',
            r'\bПрежде всего\b': 'Сначала',
            r'\bНо не в последнюю очередь\b': 'И в заключение',
            r'\bВ целом\b': 'Обычно',
            r'\bВ заключение\b': 'В конце',
            r'\bПодводя итоги\b': 'В итоге',
            r'\bВ результате\b': 'Поэтому',
            r'\bТаким образом\b': 'Значит',
            r'\bТем не менее\b': 'Но',
            r'\bС другой стороны\b': 'А вот',
            r'\bБолее того\b': 'И',
            r'\bКроме того\b': 'Также',
            r'\bВ дополнение\b': 'Еще',
            r'\bВо-первых\b': 'Сначала',
            r'\bВо-вторых\b': 'Потом',
            r'\bВ-третьих\b': 'И наконец',
        }
        
        for old_pattern, new_replacement in russian_replacements.items():
            fixed_text = re.sub(old_pattern, new_replacement, fixed_text, flags=re.IGNORECASE)
        
        return fixed_text


def save_report(results: Dict, output_path: str):
    """
    Save analysis results to a file.
    
    Args:
        results: Analysis results
        output_path: Path to save the report
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"Отчет сохранен в: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='Обнаружение текста, сгенерированного ИИ, в DOCX файлах и предложения по улучшению.'
    )
    parser.add_argument('input_file', help='Путь к входному DOCX файлу')
    parser.add_argument('-o', '--output', help='Путь для сохранения отчета анализа (формат JSON)')
    parser.add_argument('--fix', action='store_true', help='Создать исправленную версию документа')
    parser.add_argument('--fix-output', help='Путь для сохранения исправленного документа (требуется --fix)')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_file):
        print(f"Ошибка: Файл не существует: {args.input_file}")
        sys.exit(1)
    
    detector = AITextDetector()
    
    # Process the document
    results = detector.process_docx_file(args.input_file)
    
    if 'error' in results:
        print(f"Ошибка: {results['error']}")
        sys.exit(1)
    
    # Print results
    print("\n" + "="*60)
    print("РЕЗУЛЬТАТЫ ОБНАРУЖЕНИЯ ТЕКСТА ИИ")
    print("="*60)
    print(f"Файл: {results['file_path']}")
    print(f"Длина текста: {results['analysis']['total_chars']} символов, {results['analysis']['total_words']} слов")
    print(f"Оценка вероятности ИИ: {results['analysis']['ai_probability']}%")
    
    if results['analysis']['issues']:
        print("\nНайденные проблемы:")
        for issue in results['analysis']['issues']:
            print(f"  - {issue['type']}: {issue['count']} вхождений")
            if issue['examples']:
                print(f"    Примеры: {', '.join(issue['examples'][:3])}")
    else:
        print("\nКрупных шаблонов ИИ не обнаружено.")
    
    if results['suggestions']:
        print(f"\nПредложения ({len(results['suggestions'])}):")
        for i, suggestion in enumerate(results['suggestions'], 1):
            print(f"  {i}. {suggestion['issue']}")
            print(f"     {suggestion['suggestion']}")
    else:
        print("\nНет предложений по улучшению.")
    
    print("\n" + "="*60)
    
    # Save report if requested
    if args.output:
        save_report(results, args.output)
    
    # Generate fixed document if requested
    if args.fix:
        if not args.fix_output:
            print("Ошибка: --fix-output требуется при использовании --fix")
            sys.exit(1)
        
        # Extract text and apply fixes
        original_doc = Document(args.input_file)
        detector = AITextDetector()
        
        # Process paragraphs
        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():
                paragraph.text = detector.fix_text(paragraph.text)
        
        # Process tables
        for table in original_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = detector.fix_text(cell.text)
        
        # Save the fixed document
        original_doc.save(args.fix_output)
        print(f"Исправленный документ сохранен в: {args.fix_output}")


if __name__ == "__main__":
    main()